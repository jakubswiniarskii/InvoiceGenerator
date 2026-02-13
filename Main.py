import pandas as pd
from docx import Document
import os
import tempfile
import time
import comtypes.client
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import load_workbook
def load_processed_documents():
    if os.path.exists(progress_file):
        with open(progress_file, 'r') as file:
            return set(file.read().splitlines())
    return set()

def save_processed_document(document_number):
    with open(progress_file, 'a') as file:
        file.write(f"{document_number}\n")

def fill_invoice_and_generate_pdf(template_path, output_path, data ):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_vertical_align(cell)
                replace_text_in_cell(cell, data)
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_word:
        temp_word_path = temp_word.name
        doc.save(temp_word_path)
    try:
        word_to_pdf(temp_word_path, output_path)
    finally:
        os.remove(temp_word_path)


def replace_text_in_paragraph(paragraph, placeholders):
    full_text = ''.join(run.text for run in paragraph.runs)
    for key, value in placeholders.items():
        full_text = full_text.replace(key, str(value))
    if full_text != ''.join(run.text for run in paragraph.runs):
        for i, run in enumerate(paragraph.runs):
            if i == 0:
                run.text = full_text
            else:
                run.text = ''


def replace_text_in_cell(cell, placeholders):
    for paragraph in cell.paragraphs:
        replace_text_in_paragraph(paragraph, placeholders)


def cell_vertical_align(cell):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)


def word_to_pdf(word_file, pdf_file):
    wdFormatPDF = 17
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(word_file)
    doc.SaveAs(pdf_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

def format_date(date):
    return date.strftime('%Y-%m-%d') if pd.notnull(date) else ''

# Ścieżki plików
excel_file = r'C:\Users\JakubŚwiniarski\PycharmProjects\ProjektFV\Book2.xlsx'
template_path = r'C:\Users\JakubŚwiniarski\PycharmProjects\ProjektFV\Invoice_v2.docx'
output_folder = r'C:\Users\JakubŚwiniarski\PycharmProjects\ProjektFV\wygenerowaneFV'
debit_note_template = r'C:\Users\JakubŚwiniarski\PycharmProjects\ProjektFV\DebitNote.docx'
correction_invoice_template = r'C:\Users\JakubŚwiniarski\PycharmProjects\ProjektFV\CorrectionInvoice_v3.docx'
progress_file = r'C:\Users\JakubŚwiniarski\PycharmProjects\ProjektFV\progress.txt'
parking_template = r'C:\Users\JakubŚwiniarski\PycharmProjects\ProjektFV\parking.docx'
parking_correction_template = r'C:\Users\JakubŚwiniarski\PycharmProjects\ProjektFV\parking_kor.docx'
other_template = r'C:\Users\JakubŚwiniarski\PycharmProjects\ProjektFV\pozostałe.docx'
other_correction_template = r'C:\Users\JakubŚwiniarski\PycharmProjects\ProjektFV\pozostałe_kor.docx'

# Wczytanie danych excel
data = pd.read_excel(excel_file)
data.columns = data.columns.str.strip()
data['VAT Rate'] = data['VAT Rate'].astype(str).replace({'STANDARD': '20%'})
data['VAT Rate Numeric'] = (
    pd.to_numeric(
        data['VAT Rate'].str.strip('%'),
        errors='coerce'
    )
    .fillna(0)
    / 100
)
numeric_cols = [
    'Net Sales', 'Vat Amount', 'Gross Sales',
    'item Price', 'Quantity', 'Net Sales',
    'Quantity', 'Net Sales',
    'Vat Amount', 'Gross Sales'
]
for col in numeric_cols:
    if col in data.columns:
        data[col] = pd.to_numeric(
            data[col].astype(str).str.replace(',', ''),
            errors='coerce'
        )

processed_documents = load_processed_documents()
for _, row in data.iterrows():
    document_number = row['Document Number']
    unique_key = f"{document_number}_{row.get('Other.Corrected Invoice Number', 'N/A')}"

    if unique_key in processed_documents:
        continue

    # Generowanie zwykłej faktury
    if 'INV' in document_number:
        service_name = str(row.get("Document Number")).strip().lower()
        if service_name == "rent":
            invoice_data = {
                '{InvoiceNumber}': document_number,
                '{DateOfIssue}': format_date(row['Posted Date']),
                '{NameSurname}': row["Buyer's Name"],
                '{Street}': f"{row["Buyer's Address Line1"]} {row["Buyer's Address Line 2"]}".replace('nan','').strip(),
                '{City}': row["Buyer's City"],
                '{Postcode}': row["Buyer's Postcode"],
                '{Country}': row["Buyer's Country"],
                '{SellerName}': row["Originator Name"],
                '{SellerStreet}': row["Originator Address Line1"],
                '{SellerPostcodeAndCity}': f"{row["Originator Postcode"]} {row["Originator City"]}",
                '{SellerCountry}': row["Originator Country"],
                '{Currency}': row['Currency Code'],
                '{DueDate}': format_date(row['Due date']),
                '{NetValue}': row['Net Sales'],
                '{VatValue}': row['Vat Amount'],
                '{GrossVal}': row['Gross Sales'],
                '{ComDate}': format_date(row['Completion Date']),
                '{%}': row['VAT Rate'],
                '{Description}': row["Product External Description"],
                '{dateFrom}': format_date(row['Period From']),
                '{DateTo}': format_date(row['Period To']),
            }
            current_invoice_template = template_path
        elif service_name == "parking fee":
            invoice_data = {
                '{InvoiceNumber}':document_number,
                '{DateOfIssue}': format_date(row['Posted Date']),
                '{NameSurname}': row["Buyer's Name"],
                '{Street}': f"{row["Buyer's Address Line1"]} {row["Buyer's Address Line 2"]}".replace('nan','').strip(),
                '{City}': row["Buyer's City"],
                '{Postcode}': row["Buyer's Postcode"],
                '{Country}': row["Buyer's Country"],
                '{SellerName}': row["Originator Name"],
                '{SellerStreet}': row["Originator Address Line1"],
                '{SellerPostcodeAndCity}': f"{row['Originator Postcode']} {row['Originator City']}",
                '{SellerCountry}': row["Originator Country"],
                '{Currency}': row['Currency Code'],
                '{DueDate}': format_date(row['Due date']),
                '{NetValue}': row['Net Sales'],
                '{VatValue}': row['Vat Amount'],
                '{GrossVal}': row['Gross Sales'],
                '{ComDate}': format_date(row['Completion Date']),
                '{%}': row['VAT Rate'],
                '{Description}': row["Product External Description"],
                '{dateFrom}': format_date(row['Period From']),
                '{DateTo}': format_date(row['Period To']),
                '{UPrice}':row['item Price'],
                '{Quant}': row['Quantity'],
            }
            current_invoice_template = parking_template

        else:
            invoice_data = {
                '{InvoiceNumber}': document_number,
                '{DateOfIssue}': format_date(row['Posted Date']),
                '{NameSurname}': row["Buyer's Name"],
                '{Street}': f"{row["Buyer's Address Line1"]} {row["Buyer's Address Line 2"]}".replace('nan','').strip(),
                '{City}': row["Buyer's City"],
                '{Postcode}': row["Buyer's Postcode"],
                '{Country}': row["Buyer's Country"],
                '{SellerName}': row["Originator Name"],
                '{SellerStreet}': row["Originator Address Line1"],
                '{SellerPostcodeAndCity}': f"{row['Originator Postcode']} {row['Originator City']}",
                '{SellerCountry}': row["Originator Country"],
                '{Currency}': row['Currency Code'],
                '{DueDate}': format_date(row['Due date']),
                '{NetValue}': row['Net Sales'],
                '{VatValue}': row['Vat Amount'],
                '{GrossVal}': row['Gross Sales'],
                '{ComDate}': format_date(row['Completion Date']),
                '{%}': row['VAT Rate'],
                '{Description}': row["Product External Description"],
                '{UPrice}':row['item Price'],
                '{Quant}':row['Quantity'],
            }
            current_invoice_template = other_template

        pdf_file = os.path.join(output_folder, f"Invoice_{row['Document Number']}.pdf")

        fill_invoice_and_generate_pdf(current_invoice_template, pdf_file, invoice_data)

        print(f"Wygenerowano zwykłą fakturę: {pdf_file}")
        save_processed_document(unique_key)


    elif 'DN' in document_number:
        debit_note_data = {
            'XXX/07/22/DN': document_number,
            '01-07-2022': format_date(row['Other.Issue Date']),
            'Firstname Lastname': row["Other.Buyer's Name"],
            'Address line 1': row["Other.Buyer's Address Line1"],
            'Address line 2': f"{row["Other.Buyer's Postcode"]} {row["Other.Buyer's City"]} {row["Other.Buyer's Country"]}",
            'PLN': row['Other.Currency Code'],
            '05-07-2022': format_date(row['Other.Due date']),
            """Kaucja
Deposit""": row["Other.Product External Description"],
            '1,00': f"{row['Other.Quantity']:.2f}",
            '1000,00': f"{row['Other.Net Sales']:.2f}",
            '1000,00': f"{row['Other.Gross Sales']:.2f}",
            '1000,00': row['Other.Gross Sales'],
        }

        pdf_file = os.path.join(output_folder, f"DebitNote_{document_number}.pdf")

        fill_invoice_and_generate_pdf(debit_note_template, pdf_file, debit_note_data)

        print(f"Wygenerowano notę debetową: {pdf_file}")
        save_processed_document(unique_key)


# Obsługa faktur korekt
correction_invoices = data[data['Document Number'].str.contains('CI', na=False)]
grouped_corrections = correction_invoices.groupby('Corrected Invoice Number')

for original_invoice_number, group in grouped_corrections:
    group = group.sort_values(by='index')  # Sortowanie korekt w grupie po indeksie
    previous_correction_data = None

    for i, (_, row) in enumerate(group.iterrows()):
        if i == 0:
            reference_invoice = data[data['Document Number'] == original_invoice_number]
            if reference_invoice.empty:
                print(f"Nie znaleziono faktury pierwotnej dla {original_invoice_number}, pomijanie...")
                break

            reference_invoice = reference_invoice.iloc[0]
            reference_net_sales = reference_invoice['Net Sales']
            reference_vat_amount = reference_invoice['Vat Amount']
            reference_gross_sales = reference_invoice['Gross Sales']
            reference_vat_rate = reference_invoice['VAT Rate']
            reference_date_from = format_date(reference_invoice['Period From'])
            reference_date_to = format_date(reference_invoice['Period To'])
            reference_issue_date = format_date(reference_invoice['Posted Date'])
            reference_description = reference_invoice["Product External Description"]
            reference_uprice = reference_invoice["item Price"]
            reference_quant = reference_invoice["Quantity"]
        else:
            if previous_correction_data is None:
                print(f"Błąd: Brak danych poprzedniej korekty dla {original_invoice_number}.")
                break

            reference_net_sales = previous_correction_data['corrected_net_sales']
            reference_vat_amount = previous_correction_data['corrected_vat_amount']
            reference_gross_sales = previous_correction_data['corrected_gross_sales']
            reference_vat_rate = previous_correction_data['corrected_vat_rate']
            reference_date_from = previous_correction_data['corrected_date_from']
            reference_date_to = previous_correction_data['corrected_date_to']
            reference_description = previous_correction_data['description']
            reference_uprice = previous_correction_data['uprice']
            reference_quant = previous_correction_data['quant']

        # Obliczenia wartości netto, VAT i brutto po korekcie
        corrected_net_sales = reference_net_sales + row['Net Sales']
        vat_rate = float(str(row['VAT Rate']).strip('%')) / 100
        corrected_vat_amount = corrected_net_sales * vat_rate
        corrected_gross_sales = corrected_net_sales + corrected_vat_amount

        # Obliczenie różnic względem referencyjnej faktury
        difference_net_sales = corrected_net_sales - reference_net_sales
        difference_vat_amount = corrected_vat_amount - reference_vat_amount
        difference_gross_sales = corrected_gross_sales - reference_gross_sales

        service_name = str(row.get("Service Name")).strip().lower()
        if service_name == "rent":
            current_correction_template = correction_invoice_template
            correction_data = {
                '{InvoiceNumber}': row['Document Number'],
                '{RInvoiceNumber}': row['Corrected Invoice Number'],
                '{DateOfIssue}': format_date(row['Posted Date']),
                '{RDateOfIssue}': reference_issue_date,
                '{Currency}': row["Currency Code"],
                '{NameSurname}': row["Buyer's Name"],
                "{Street}": f"{row["Buyer's Address Line1"]} {row["Buyer's Address Line 2"]}".replace('nan','').strip(),
                '{City}':row["Buyer's City"],
                '{Postcode}': row["Buyer's Postcode"],
                '{Country}': row["Buyer's Country"],
                '{SellerName}': row["Originator Name"],
                '{SellerStreet}': row["Originator Address Line1"],
                '{SellerPostcodeCity}': f"{row["Originator Postcode"]} {row["Originator City"]}",
                '{SellerCountry}': row["Originator Country"],
                '{Description}': row["Product External Description"],
                '{RNetVal}': f"{reference_invoice['Net Sales']:.2f}",
                '{RDescription}': reference_description,
                '{R%}': reference_vat_rate,
                '{RFrom}': reference_date_from,
                '{RTo}': reference_date_to,
                '{DateFrom}': format_date(row["Period From"]),
                '{DateTo}': format_date(row["Period To"]),
                '{RVatVal}': f"{reference_invoice['Vat Amount']:.2f}",
                '{RGrossVal}': f"{reference_invoice['Gross Sales']:.2f}",
                '{CNetVal}': f"{corrected_net_sales:.2f}",
                '{CVatVal}': f"{corrected_vat_amount:.2f}",
                '{CGrossV}': f"{corrected_gross_sales:.2f}",
                '{%}': row['VAT Rate'],
                '{ComDate}': format_date(row["Completion Date"]),
                '{DnetV}': f"{difference_net_sales:+.2f}",
                '{DVatV}': f"{difference_vat_amount:+.2f}",
                "{DGross}": f"{difference_gross_sales:+.2f}",
                '{Reason}': "Change of contract terms / Zmiana warunków umowy"
            }
        elif service_name == "parking fee":

            quantity_change  = row['Quantity']

            if row['New Net Sales'] > 0:
                new_quantity = reference_quant + quantity_change
            elif row['New Net Sales'] < 0:
                new_quantity = reference_quant - quantity_change

            current_correction_template = parking_correction_template
            correction_data = {
                '{InvoiceNumber}': row['Document Number'],
                '{RInvoiceNumber}': row['Corrected Invoice Number'],
                '{DateOfIssue}': format_date(row['Posted Date']),
                '{RDateOfIssue}': reference_issue_date,
                '{Currency}': row["Currency Code"],
                '{NameSurname}': row["Buyer's Name"],
                '{Street}': f"{row["Buyer's Address Line1"]} {row["Buyer's Address Line 2"]}".replace('nan','').strip(),
                '{City}': row["Buyer's City"],
                '{Postcode}': row["Buyer's Postcode"],
                '{Country}': row["Buyer's Country"],
                '{SellerName}': row["Originator Name"],
                '{SellerStreet}': row["Originator Address Line1"],
                '{SellerPostcodeCity}': f"{row['Originator Postcode']} {row['Originator City']}",
                '{SellerCountry}': row["Originator Country"],
                '{Description}': row["Product External Description"],
                '{RNetVal}': f"{reference_invoice['Net Sales']:.2f}",
                '{RDescription}': reference_description,
                '{RUPrice}': abs(reference_uprice),
                '{RQuant}': reference_quant,
                '{R%}': reference_vat_rate,
                '{RFrom}': reference_date_from,
                '{RTo}': reference_date_to,
                '{DateFrom}': format_date(row["Period From"]),
                '{DateTo}': format_date(row["Period To"]),
                '{RVatVal}': f"{reference_invoice['Vat Amount']:.2f}",
                '{RGrossVal}': f"{reference_invoice['Gross Sales']:.2f}",
                '{CNetVal}': f"{corrected_net_sales:.2f}",
                '{CVatVal}': f"{corrected_vat_amount:.2f}",
                '{CGrossV}': f"{corrected_gross_sales:.2f}",
                '{%}': row['VAT Rate'],
                '{UPrice}':abs(row['item Price']),
                '{Quant}': new_quantity,
                '{ComDate}': format_date(row["Completion Date"]),
                '{DnetV}': f"{difference_net_sales:+.2f}",
                '{DVatV}': f"{difference_vat_amount:+.2f}",
                '{DGross}': f"{difference_gross_sales:+.2f}",
                '{Reason}': "Change of contract terms / Zmiana warunków umowy"
            }
        else:
            quantity_change = row['Quantity']

            if row['Net Sales'] > 0:
                new_quantity = reference_quant + quantity_change
            elif row['Net Sales'] < 0:
                new_quantity = reference_quant - quantity_change
            current_correction_template = other_correction_template
            correction_data = {
                '{InvoiceNumber}': row['Document Number'],
                '{RInvoiceNumber}': row['Corrected Invoice Number'],
                '{DateOfIssue}': format_date(row['Posted Date']),
                '{RDateOfIssue}': reference_issue_date,
                '{Currency}': row["Currency Code"],
                '{NameSurname}': row["Buyer's Name"],
                '{Street}': f"{row["Buyer's Address Line1"]} {row["Buyer's Address Line 2"]}".replace('nan','').strip(),
                '{City}': row["Buyer's City"],
                '{Postcode}': row["Buyer's Postcode"],
                '{Country}': row["Buyer's Country"],
                '{SellerName}': row["Originator Name"],
                '{SellerStreet}': row["Originator Address Line1"],
                '{SellerPostcodeCity}': f"{row['Originator Postcode']} {row['Originator City']}",
                '{SellerCountry}': row["Originator Country"],
                '{Description}': row["Product External Description"],
                '{RNetVal}': f"{reference_invoice['Net Sales']:.2f}",
                '{RDescription}': reference_description,
                '{RUPrice}': abs(reference_uprice),
                '{RQuant}': reference_quant,
                '{R%}': reference_vat_rate,
                '{RVatVal}': f"{reference_invoice['Vat Amount']:.2f}",
                '{RGrossVal}': f"{reference_invoice['Gross Sales']:.2f}",
                '{CNetVal}': f"{corrected_net_sales:.2f}",
                '{CVatVal}': f"{corrected_vat_amount:.2f}",
                '{CGrossV}': f"{corrected_gross_sales:.2f}",
                '{%}': row['VAT Rate'],
                '{UPrice}':abs(row['item Price']),
                '{Quant}': new_quantity,
                '{ComDate}': format_date(row["Completion Date"]),
                '{DnetV}': f"{difference_net_sales:+.2f}",
                '{DVatV}': f"{difference_vat_amount:+.2f}",
                '{DGross}': f"{difference_gross_sales:+.2f}",
                '{Reason}': "Change of contract terms / Zmiana warunków umowy"
            }


        pdf_file = os.path.join(output_folder, f"CorrectionInvoice_{original_invoice_number}_{row['Document Number']}.pdf")
        fill_invoice_and_generate_pdf(current_correction_template, pdf_file, correction_data)

        print(f"Wygenerowano fakturę korektę: {pdf_file}")
        save_processed_document(unique_key)


        previous_correction_data = {
            'corrected_net_sales': corrected_net_sales,
            'corrected_vat_amount': corrected_vat_amount,
            'corrected_gross_sales': corrected_gross_sales,
            'corrected_vat_rate': reference_vat_rate,
            'corrected_date_from': reference_date_from,
            'corrected_date_to': reference_date_to,
            'description' : reference_description,
            'uprice' : reference_uprice,
            'quant' : reference_quant,
        }
        reference_invoice = {
            'Document Number': row['Document Number'],
            'Net Sales': corrected_net_sales,
            'Vat Amount': corrected_vat_amount,
            'Gross Sales': corrected_gross_sales,
            'corrected_vat_rate': reference_vat_rate,
            'corrected_date_from': reference_date_from,
            'corrected_date_to': reference_date_to,
            'description': reference_description,
            'uprice': reference_uprice,
            'quant': reference_quant,

        }








